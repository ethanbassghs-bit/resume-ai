
# resume_ai_builder.py

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import re

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def extract_end_date(text):
    match = re.search(r"(?:–|-|—)\s*([A-Za-z]+\s+\d{4}|Present)", text)
    if match:
        date_str = match.group(1)
        if "Present" in date_str:
            return datetime.today()
        try:
            return datetime.strptime(date_str, "%B %Y")
        except ValueError:
            pass
    return datetime.min

def format_resume(text_sections):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Garamond"
    font.size = Pt(11)

    # Header name
    header = doc.add_paragraph()
    run = header.add_run("Ethan S. Bass")
    run.bold = True
    run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    table = doc.add_table(rows=2, cols=3)
    widths = [Inches(2.6), Inches(2.7), Inches(2.6)]
    for i, width in enumerate(widths):
        table.columns[i].width = width

    row1 = table.rows[0].cells
    row1[0].text = "1091 McLynn Avenue"
    row1[1].text = "(404) 435-1520 | ethanbassghs@gmail.com"
    row1[2].text = "170 River Road"

    row2 = table.rows[1].cells
    row2[0].text = "Atlanta, GA 30306"
    row2[1].text = ""
    row2[2].text = "Athens, GA 30605"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for r in p.runs:
                    r.font.name = "Garamond"
                    r.font.size = Pt(11)

    doc.add_paragraph("")

    for section_title, content in text_sections.items():
        doc.add_paragraph("")
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title.upper())
        heading_run.bold = True
        heading_run.font.name = "Garamond"
        heading_run.font.size = Pt(11)

        sorted_content = sorted(content, key=extract_end_date, reverse=True)

        i = 0
        while i < len(sorted_content):
            line = sorted_content[i]
            if line and line[0].isupper() and "," in line and len(line.split()) < 20:
                job_title_line = line
                i += 1
                bullet_points = []
                while i < len(sorted_content) and not (sorted_content[i] and sorted_content[i][0].isupper() and "," in sorted_content[i]):
                    bullet_points.append(sorted_content[i])
                    i += 1
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(job_title_line.strip())
                job_run.bold = True
                job_run.font.name = "Garamond"
                job_run.font.size = Pt(11)
                for bp in bullet_points:
                    bullet = doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet.add_run(bp.strip("• "))
                    bullet_run.font.name = "Garamond"
                    bullet_run.font.size = Pt(11)
            else:
                i += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def enhance_text(input_text):
    prompt = (
        "Improve the following resume bullet points to be stronger, more action-oriented, and quantifiable. "
        "Keep them professional and concise.\n\n" + input_text
    )
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# --- Streamlit App UI ---
st.set_page_config(page_title="Resume AI", page_icon="📄", layout="centered")
st.title("📄 Apollo Society Resume Builder")
st.markdown("""
Build and enhance resumes in seconds. Upload your draft or paste resume bullets, and select the **Apollo Society** template.
""")

format_option = st.selectbox("Choose resume format:", ["Apollo Society", "Custom (coming soon)"])
input_method = st.radio("Input method:", ["Paste text", "Upload .docx"])

raw_text = ""
if input_method == "Paste text":
    raw_text = st.text_area("Paste your resume bullets or content here:", height=300)
elif input_method == "Upload .docx":
    uploaded_file = st.file_uploader("Upload Word Document:", type=["docx"])
    if uploaded_file:
        doc = Document(uploaded_file)
        raw_text = "\n".join([para.text for para in doc.paragraphs])

if raw_text:
    if st.button("✨ Enhance Bullets with AI"):
        improved_text = enhance_text(raw_text)
        st.subheader("Enhanced Resume Bullets:")
        st.text(improved_text)

    if st.button("📁 Format Resume to Apollo Society Template"):
        sections = {}
        current_section = "Other"
        sections[current_section] = []
        for line in raw_text.splitlines():
            if line.strip().isupper() and len(line.strip()) < 40:
                current_section = line.strip()
                sections[current_section] = []
            elif line.strip():
                sections[current_section].append(line.strip("• "))

        resume_docx = format_resume(sections)
        st.download_button(
            label=f"📥 Download {format_option} Resume",
            data=resume_docx,
            file_name=f"{format_option.replace(' ', '_')}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
