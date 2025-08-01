# resume_ai_builder.py

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import re

# Configure OpenAI with new SDK
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# --- Helper: Extract end date for sorting ---
def extract_end_date(text):
    match = re.search(r"–\s*([A-Za-z]+\s+\d{4}|Present)", text)
    if match:
        date_str = match.group(1)
        if "Present" in date_str:
            return datetime.today()
        try:
            return datetime.strptime(date_str, "%B %Y")
        except ValueError:
            pass
    return datetime.min  # Default if no valid date

# --- Helper: Format Resume in Apollo Society Template ---
def format_resume(text_sections):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11)

    header = doc.add_paragraph()
    run = header.add_run("Ethan S. Bass")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    row_cells = table.rows[0].cells
    cell_left = row_cells[0].paragraphs[0].add_run("1091 McLynn Avenue")
    cell_left.font.name = 'Garamond'
    cell_left.font.size = Pt(11)
    cell_right = row_cells[1].paragraphs[0].add_run("(404) 435-1520 | ethanbassghs@gmail.com | Athens, GA 30605")
    cell_right.font.name = 'Garamond'
    cell_right.font.size = Pt(11)

    for section_title, content in text_sections.items():
        doc.add_paragraph("\n")
        heading = doc.add_paragraph()
        run = heading.add_run(section_title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Garamond'

        sorted_content = sorted(content, key=extract_end_date, reverse=True)

        for bullet in sorted_content:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1

            parts = [part.strip() for part in bullet.split(',')]
            if len(parts) >= 2:
                org = parts[0]
                role = parts[1]
                rest = ", ".join(parts[2:]) if len(parts) > 2 else ""

                r = p.add_run(org)
                r.bold = True
                r.font.name = 'Garamond'
                r.font.size = Pt(11)

                p.add_run(f", {role}").font.size = Pt(11)
                if rest:
                    p.add_run(f", {rest}").font.size = Pt(11)
            else:
                r = p.add_run(bullet)
                r.font.name = 'Garamond'
                r.font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Helper: Enhance Resume Text with AI ---
def enhance_text(input_text):
    prompt = (
        "Improve the following resume bullet points to be stronger, more action-oriented, and quantifiable."
        " Keep them professional and concise:\n\n" + input_text
    )
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# --- Streamlit Interface ---
st.title("📄 Resume Formatter & Enhancer")
st.write("Upload your resume or paste bullets below to format & improve them.")

format_option = st.selectbox("Choose resume format:", ["Apollo Society", "Custom (coming soon)"])
input_method = st.radio("Choose input method:", ["Paste text", "Upload .docx"])

raw_text = ""
if input_method == "Paste text":
    raw_text = st.text_area("Paste your resume bullets or text:", height=300)
elif input_method == "Upload .docx":
    uploaded_file = st.file_uploader("Upload Word Document:", type=["docx"])
    if uploaded_file:
        doc = Document(uploaded_file)
        raw_text = "\n".join([para.text for para in doc.paragraphs])

if raw_text:
    if st.button("✨ Enhance Text with AI"):
        improved_text = enhance_text(raw_text)
        st.subheader("Enhanced Resume Bullets:")
        st.text(improved_text)

    if st.button("📁 Format into Resume Template"):
        sections = {}
        current_section = "Other"
        sections[current_section] = []
        for line in raw_text.splitlines():
            if line.strip().isupper() and len(line.strip()) < 40:
                current_section = line.strip()
                sections[current_section] = []
            elif line.strip():
                sections[current_section].append(line.strip("• "))

        resume_docx = format_resume(sections)
        st.download_button(
            label=f"📥 Download {format_option} Resume",
            data=resume_docx,
            file_name=f"{format_option.replace(' ', '_')}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
